"""
generate_test_docx.py
Generates a comprehensive DOCX test document covering all standard
word-processor features needed in a multi-format publishing pipeline.
"""

import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree
import copy

# ──────────────────────────────────────────────
# Helper utilities
# ──────────────────────────────────────────────

def add_horizontal_rule(doc):
    """Insert a paragraph-level horizontal rule using w:pBdr/w:bottom."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def set_cell_shading(cell, fill_hex):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def set_cell_borders(cell, **kwargs):
    """Set individual cell borders (top/bottom/left/right)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if side in kwargs:
            tag = OxmlElement(f"w:{side}")
            tag.set(qn("w:val"), kwargs[side].get("val", "single"))
            tag.set(qn("w:sz"), str(kwargs[side].get("sz", 4)))
            tag.set(qn("w:color"), kwargs[side].get("color", "auto"))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_bookmark(paragraph, bookmark_id, name):
    """Wrap paragraph content in a bookmark."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")
    bookmarkStart = OxmlElement("w:bookmarkStart")
    bookmarkStart.set(qn("w:id"), str(bookmark_id))
    bookmarkStart.set(qn("w:name"), name)
    bookmarkEnd = OxmlElement("w:bookmarkEnd")
    bookmarkEnd.set(qn("w:id"), str(bookmark_id))
    run._r.addprevious(bookmarkStart)
    run._r.addnext(bookmarkEnd)


def add_hyperlink(paragraph, url, text, color="0563C1", underline=True):
    """Add a clickable hyperlink run to an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_footnote(paragraph, note_text):
    """
    Append a footnote reference to the last run in *paragraph*
    and register the footnote body in the document's footnotes part.
    Uses low-level XML because python-docx has no footnote API.
    """
    doc = paragraph.part.document

    # ── locate / build footnotes part ────────────────────────────────────────
    fn_part = None
    for rel in doc.part.rels.values():
        if "footnotes" in rel.reltype:
            fn_part = rel._target
            break

    if fn_part is None:
        # Create the footnotes part from scratch
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        fn_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:footnotes xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"'
            ' xmlns:cx="http://schemas.microsoft.com/office/drawing/2014/chartex"'
            ' xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"'
            ' xmlns:o="urn:schemas-microsoft-com:office:office"'
            ' xmlns:oel="http://schemas.microsoft.com/office/2019/extlst"'
            ' xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
            ' xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'
            ' xmlns:v="urn:schemas-microsoft-com:vml"'
            ' xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"'
            ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
            ' xmlns:w10="urn:schemas-microsoft-com:office:word"'
            ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"'
            ' xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"'
            ' xmlns:w16cex="http://schemas.microsoft.com/office/word/2018/wordml/cex"'
            ' xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid"'
            ' xmlns:w16="http://schemas.microsoft.com/office/word/2018/wordml"'
            ' xmlns:w16sdtdh="http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash"'
            ' xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex"'
            ' xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"'
            ' xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"'
            ' xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"'
            ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
            # separator footnote (id=-1)
            '<w:footnote w:type="separator" w:id="-1">'
            '<w:p><w:r><w:separator/></w:r></w:p>'
            '</w:footnote>'
            # continuation separator (id=0)
            '<w:footnote w:type="continuationSeparator" w:id="0">'
            '<w:p><w:r><w:continuationSeparator/></w:r></w:p>'
            '</w:footnote>'
            '</w:footnotes>'
        )
        fn_part = Part(
            PackURI("/word/footnotes.xml"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
            fn_xml.encode("utf-8"),
            doc.part.package,
        )
        doc.part.relate_to(
            fn_part,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
        )

    # ── determine next footnote id ────────────────────────────────────────────
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    fn_root = etree.fromstring(fn_part.blob)
    existing_ids = [
        int(fn.get(f"{{{W}}}id"))
        for fn in fn_root.findall(f"{{{W}}}footnote")
        if fn.get(f"{{{W}}}id") not in ("-1", "0")
    ]
    fn_id = max(existing_ids, default=0) + 1

    # ── build the footnote element ────────────────────────────────────────────
    fn_el = etree.SubElement(fn_root, f"{{{W}}}footnote")
    fn_el.set(f"{{{W}}}id", str(fn_id))
    fn_p = etree.SubElement(fn_el, f"{{{W}}}p")
    fn_pPr = etree.SubElement(fn_p, f"{{{W}}}pPr")
    fn_pStyle = etree.SubElement(fn_pPr, f"{{{W}}}pStyle")
    fn_pStyle.set(f"{{{W}}}val", "Footnotetext")
    fn_r0 = etree.SubElement(fn_p, f"{{{W}}}r")
    fn_rPr = etree.SubElement(fn_r0, f"{{{W}}}rPr")
    fn_rStyle = etree.SubElement(fn_rPr, f"{{{W}}}rStyle")
    fn_rStyle.set(f"{{{W}}}val", "FootnoteReference")
    fn_ref = etree.SubElement(fn_r0, f"{{{W}}}footnoteRef")
    fn_r1 = etree.SubElement(fn_p, f"{{{W}}}r")
    fn_t = etree.SubElement(fn_r1, f"{{{W}}}t")
    fn_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    fn_t.text = " " + note_text

    fn_part._blob = etree.tostring(fn_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    # ── insert footnote reference run into the paragraph ──────────────────────
    ref_run = OxmlElement("w:r")
    ref_rPr = OxmlElement("w:rPr")
    ref_rStyle = OxmlElement("w:rStyle")
    ref_rStyle.set(qn("w:val"), "FootnoteReference")
    ref_rPr.append(ref_rStyle)
    ref_run.append(ref_rPr)
    ref_ref = OxmlElement("w:footnoteReference")
    ref_ref.set(qn("w:id"), str(fn_id))
    ref_run.append(ref_ref)
    paragraph._p.append(ref_run)


def add_toc_field(doc):
    """Insert a TOC field code paragraph."""
    p = doc.add_paragraph()
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")
    run = p.add_run()
    run._r.append(fldChar_begin)
    run2 = p.add_run()
    run2._r.append(instrText)
    run3 = p.add_run()
    run3._r.append(fldChar_end)
    p.style = "TOC Heading"
    return p


def add_page_number_to_footer(section):
    """Add centered 'Page X of Y' to the footer of a section."""
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()

    def _field_run(para, instr):
        r = OxmlElement("w:r")
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        r.append(begin)
        para._p.append(r)
        r2 = OxmlElement("w:r")
        it = OxmlElement("w:instrText")
        it.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        it.text = f" {instr} "
        r2.append(it)
        para._p.append(r2)
        r3 = OxmlElement("w:r")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        r3.append(end)
        para._p.append(r3)

    run = para.add_run("Page ")
    _field_run(para, "PAGE")
    para.add_run(" of ")
    _field_run(para, "NUMPAGES")


def add_math_equation(doc, latex_like):
    """Insert a simple OMML equation (x² + y² = z²)."""
    # Build minimal OMML for "x² + y² = z²"
    omml = (
        '<m:oMathPara xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">'
        "<m:oMath>"
        "<m:r><m:t>x</m:t></m:r>"
        "<m:sSup>"
        "<m:sSupPr><m:ctrlPr/></m:sSupPr>"
        "<m:e><m:r><m:t>2</m:t></m:r></m:e>"
        "<m:sup><m:r><m:t>&#160;</m:t></m:r></m:sup>"
        "</m:sSup>"
        "<m:r><m:t> + y</m:t></m:r>"
        "<m:sSup>"
        "<m:sSupPr><m:ctrlPr/></m:sSupPr>"
        "<m:e><m:r><m:t>2</m:t></m:r></m:e>"
        "<m:sup><m:r><m:t>&#160;</m:t></m:r></m:sup>"
        "</m:sSup>"
        "<m:r><m:t> = z</m:t></m:r>"
        "<m:sSup>"
        "<m:sSupPr><m:ctrlPr/></m:sSupPr>"
        "<m:e><m:r><m:t>2</m:t></m:r></m:e>"
        "<m:sup><m:r><m:t>&#160;</m:t></m:r></m:sup>"
        "</m:sSup>"
        "</m:oMath>"
        "</m:oMathPara>"
    )
    math_el = etree.fromstring(omml)
    p = doc.add_paragraph()
    p._p.append(math_el)
    return p


def apply_highlight(run, color_name):
    """Apply highlight color to a run (e.g. 'yellow', 'cyan')."""
    rPr = run._r.get_or_add_rPr()
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), color_name)
    rPr.append(highlight)


def apply_strikethrough(run):
    rPr = run._r.get_or_add_rPr()
    strike = OxmlElement("w:strike")
    rPr.append(strike)


def apply_double_strike(run):
    rPr = run._r.get_or_add_rPr()
    ds = OxmlElement("w:dstrike")
    rPr.append(ds)


def apply_small_caps(run):
    rPr = run._r.get_or_add_rPr()
    sc = OxmlElement("w:smallCaps")
    rPr.append(sc)


def apply_all_caps(run):
    rPr = run._r.get_or_add_rPr()
    caps = OxmlElement("w:caps")
    rPr.append(caps)


def set_line_spacing(paragraph, rule, value_twips=None):
    """Set paragraph line spacing."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    if rule == "exact":
        spacing.set(qn("w:line"), str(value_twips))
        spacing.set(qn("w:lineRule"), "exact")
    elif rule == "double":
        spacing.set(qn("w:line"), "480")
        spacing.set(qn("w:lineRule"), "auto")
    elif rule == "1.5":
        spacing.set(qn("w:line"), "360")
        spacing.set(qn("w:lineRule"), "auto")


def add_drop_cap(paragraph):
    """Apply a 3-line drop-cap to the first character via w:framePr."""
    # Drop caps are complex; we mark the paragraph style instead via XML
    pPr = paragraph._p.get_or_add_pPr()
    framePr = OxmlElement("w:framePr")
    framePr.set(qn("w:dropCap"), "drop")
    framePr.set(qn("w:lines"), "3")
    framePr.set(qn("w:wrap"), "around")
    framePr.set(qn("w:vAnchor"), "text")
    framePr.set(qn("w:hAnchor"), "text")
    pPr.append(framePr)


def add_comment(doc, paragraph, comment_text, author="Test Author"):
    """Add a review comment to a paragraph (basic implementation)."""
    # Comments require the comments part — simplified: just add a comment mark
    # This is a known complex operation; we include it as XML placeholder
    run = paragraph.add_run("")
    commentStart = OxmlElement("w:commentRangeStart")
    commentStart.set(qn("w:id"), "1")
    commentEnd = OxmlElement("w:commentRangeEnd")
    commentEnd.set(qn("w:id"), "1")
    commentRef = OxmlElement("w:commentReference")
    commentRef.set(qn("w:id"), "1")
    paragraph._p.insert(0, commentStart)
    paragraph._p.append(commentEnd)
    run._r.append(commentRef)


# ══════════════════════════════════════════════
# BUILD THE DOCUMENT
# ══════════════════════════════════════════════

doc = Document()

# ── Page setup ────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ── Header ────────────────────────────────────
header = section.header
hdr_para = header.paragraphs[0]
hdr_para.text = "DOCX Pipeline Test Document"
hdr_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# ── Footer with page numbers ──────────────────
add_page_number_to_footer(section)

# ══════════════════════════════════════════════
# SECTION 1 – Cover / Title block
# ══════════════════════════════════════════════

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_para.add_run("DOCX Multi-Format Pipeline\nTest Document")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle_para = doc.add_paragraph()
subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle_para.add_run("A comprehensive specimen of word-processor features\nfor conversion pipeline validation")
r.italic = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x44, 0x47, 0x2A)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Version 1.0  |  2026-05-20  |  CC BY 4.0").font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════

doc.add_heading("Contents", level=1)
toc_p = doc.add_paragraph()
toc_p.style = "Normal"
fldChar1 = OxmlElement("w:fldChar")
fldChar1.set(qn("w:fldCharType"), "begin")
instrText = OxmlElement("w:instrText")
instrText.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
fldChar2 = OxmlElement("w:fldChar")
fldChar2.set(qn("w:fldCharType"), "separate")
fldChar3 = OxmlElement("w:fldChar")
fldChar3.set(qn("w:fldCharType"), "end")
r0 = toc_p.add_run()
r0._r.append(fldChar1)
r1 = toc_p.add_run()
r1._r.append(instrText)
r2 = toc_p.add_run()
r2._r.append(fldChar2)
toc_p.add_run("(Right-click → Update Field to populate)")
r3 = toc_p.add_run()
r3._r.append(fldChar3)

doc.add_page_break()

# ══════════════════════════════════════════════
# SECTION 2 – Headings (H1–H4)
# ══════════════════════════════════════════════

h1 = doc.add_heading("1  Typography & Inline Formatting", level=1)
add_bookmark(h1, 1, "sec_typography")

doc.add_heading("1.1  Heading Level 2", level=2)
doc.add_heading("1.1.1  Heading Level 3", level=3)
doc.add_heading("1.1.1.1  Heading Level 4", level=4)

# ── Body text ─────────────────────────────────
body = doc.add_paragraph(
    "This paragraph demonstrates standard body text. "
)
body.add_run("Bold text. ").bold = True
body.add_run("Italic text. ").italic = True
bold_italic = body.add_run("Bold-italic text. ")
bold_italic.bold = True
bold_italic.italic = True
underline_run = body.add_run("Underlined text. ")
underline_run.underline = True
strike_run = body.add_run("Strikethrough. ")
apply_strikethrough(strike_run)
dstrike_run = body.add_run("Double-strikethrough. ")
apply_double_strike(dstrike_run)

body2 = doc.add_paragraph()
body2.add_run("Small-caps run. ")
apply_small_caps(body2.runs[-1])
body2.add_run("All-caps run. ")
apply_all_caps(body2.runs[-1])

# ── Superscript / Subscript ───────────────────
ss_para = doc.add_paragraph("Chemical formula: H")
sub_run = ss_para.add_run("2")
sub_run.font.subscript = True
ss_para.add_run("O  |  Footnote reference")
fn_anchor = ss_para.add_run("1")
fn_anchor.font.superscript = True
ss_para.add_run(" (see footnote).")
add_footnote(ss_para, "This is footnote 1 — water is H₂O, the most abundant compound on Earth.")

# ── Font colours ──────────────────────────────
colour_para = doc.add_paragraph("Colour palette: ")
colours = [
    ("Red",   RGBColor(0xC0, 0x00, 0x00)),
    ("Green", RGBColor(0x00, 0x70, 0x00)),
    ("Blue",  RGBColor(0x00, 0x00, 0xFF)),
    ("Teal",  RGBColor(0x00, 0x80, 0x80)),
]
for label, rgb in colours:
    r = colour_para.add_run(label + "  ")
    r.font.color.rgb = rgb
    r.bold = True

# ── Highlights ────────────────────────────────
hl_para = doc.add_paragraph("Highlights: ")
for hl_text, hl_color in [
    ("Yellow", "yellow"), ("Cyan", "cyan"), ("Green", "green"),
    ("Magenta", "magenta"), ("DarkRed", "darkRed"),
]:
    run = hl_para.add_run(hl_text + " ")
    apply_highlight(run, hl_color)

# ── Font sizes ────────────────────────────────
size_para = doc.add_paragraph("Font sizes: ")
for pts in [8, 9, 10, 11, 12, 14, 16, 18, 24]:
    r = size_para.add_run(f"{pts}pt ")
    r.font.size = Pt(pts)

# ── Fonts / typefaces ─────────────────────────
font_para = doc.add_paragraph("Typefaces: ")
for face in ["Arial", "Times New Roman", "Courier New", "Georgia", "Calibri"]:
    r = font_para.add_run(face + "  ")
    r.font.name = face

# ── Spacing ───────────────────────────────────
doc.add_heading("1.2  Line & Paragraph Spacing", level=2)
p_single = doc.add_paragraph("Single line spacing — Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor.")
p_1half = doc.add_paragraph("1.5 line spacing — Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor.")
set_line_spacing(p_1half, "1.5")
p_double = doc.add_paragraph("Double line spacing — Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor.")
set_line_spacing(p_double, "double")

# ── Alignment ─────────────────────────────────
doc.add_heading("1.3  Paragraph Alignment", level=2)
for align, label in [
    (WD_ALIGN_PARAGRAPH.LEFT,     "Left-aligned text."),
    (WD_ALIGN_PARAGRAPH.CENTER,   "Centre-aligned text."),
    (WD_ALIGN_PARAGRAPH.RIGHT,    "Right-aligned text."),
    (WD_ALIGN_PARAGRAPH.JUSTIFY,  "Justified text — Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua."),
]:
    p = doc.add_paragraph(label)
    p.alignment = align

# ── Indentation ───────────────────────────────
doc.add_heading("1.4  Indentation & Block Quote", level=2)
indent_para = doc.add_paragraph(
    "This paragraph has a left indent of 1 inch and first-line indent of −0.5 in (hanging indent)."
)
indent_para.paragraph_format.left_indent   = Inches(1)
indent_para.paragraph_format.first_line_indent = Inches(-0.5)

blockquote = doc.add_paragraph(
    '"The greatest shortcoming of the human race is our inability to understand the exponential function."\n'
    "— Albert Bartlett"
)
blockquote.paragraph_format.left_indent  = Inches(0.75)
blockquote.paragraph_format.right_indent = Inches(0.75)
blockquote.paragraph_format.space_before = Pt(6)
blockquote.paragraph_format.space_after  = Pt(6)
blockquote.runs[0].italic = True

# ══════════════════════════════════════════════
# SECTION 3 – Lists
# ══════════════════════════════════════════════

doc.add_heading("2  Lists", level=1)

doc.add_heading("2.1  Unordered (Bullet) List", level=2)
for item in [
    "First bullet item",
    "Second bullet item with sub-items:",
    "Sub-item A",
    "Sub-item B",
    "Sub-sub-item i",
    "Sub-sub-item ii",
    "Third bullet item",
]:
    if "Sub-sub" in item:
        p = doc.add_paragraph(item, style="List Bullet 3")
    elif "Sub-item" in item:
        p = doc.add_paragraph(item, style="List Bullet 2")
    else:
        p = doc.add_paragraph(item, style="List Bullet")

doc.add_heading("2.2  Ordered (Numbered) List", level=2)
for item in [
    "First numbered item",
    "Second numbered item",
    "Sub-numbered item a",
    "Sub-numbered item b",
    "Third numbered item",
]:
    if "Sub-numbered" in item:
        p = doc.add_paragraph(item, style="List Number 2")
    else:
        p = doc.add_paragraph(item, style="List Number")

doc.add_heading("2.3  Definition-style List", level=2)
defs = [
    ("Term 1", "Definition of term 1 — an explanation of what it means."),
    ("Term 2", "Definition of term 2 — another explanation."),
]
for term, defn in defs:
    p = doc.add_paragraph()
    r = p.add_run(term + ":  ")
    r.bold = True
    p.add_run(defn)
    p.paragraph_format.left_indent = Inches(0.5)

# ══════════════════════════════════════════════
# SECTION 4 – Tables
# ══════════════════════════════════════════════

doc.add_heading("3  Tables", level=1)

doc.add_heading("3.1  Basic Table with Header Row", level=2)
table1 = doc.add_table(rows=5, cols=4)
table1.style = "Table Grid"
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Country", "Capital", "Population (M)", "GDP (USD tn)"]
for i, hdr in enumerate(headers):
    cell = table1.rows[0].cells[i]
    cell.text = hdr
    set_cell_shading(cell, "1F497D")
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

rows_data = [
    ["Germany",        "Berlin",    "83.2",  "4.1"],
    ["France",         "Paris",     "67.4",  "2.9"],
    ["United Kingdom", "London",    "67.2",  "3.1"],
    ["Spain",          "Madrid",    "47.4",  "1.4"],
]
for r_idx, row_data in enumerate(rows_data, start=1):
    for c_idx, val in enumerate(row_data):
        cell = table1.rows[r_idx].cells[c_idx]
        cell.text = val
        if r_idx % 2 == 0:
            set_cell_shading(cell, "DCE6F1")

doc.add_paragraph()  # spacing

doc.add_heading("3.2  Table with Merged Cells", level=2)
table2 = doc.add_table(rows=3, cols=3)
table2.style = "Table Grid"
# Merge top row across all 3 cols
table2.rows[0].cells[0].merge(table2.rows[0].cells[2])
table2.rows[0].cells[0].text = "Merged Header spanning all three columns"
set_cell_shading(table2.rows[0].cells[0], "4472C4")
table2.rows[0].cells[0].paragraphs[0].runs[0].bold = True
table2.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
table2.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
# Merge left col rows 1-2
table2.rows[1].cells[0].merge(table2.rows[2].cells[0])
table2.rows[1].cells[0].text = "Row-merged cell"
table2.rows[1].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
for r, c, txt in [(1,1,"R1 C2"), (1,2,"R1 C3"), (2,1,"R2 C2"), (2,2,"R2 C3")]:
    table2.rows[r].cells[c].text = txt

doc.add_paragraph()

# ══════════════════════════════════════════════
# SECTION 5 – Code Blocks
# ══════════════════════════════════════════════

doc.add_heading("4  Code & Pre-formatted Text", level=1)

doc.add_heading("4.1  Inline Code", level=2)
inline_p = doc.add_paragraph("Use the function ")
code_run = inline_p.add_run("calculate_p_value(data, alpha=0.05)")
code_run.font.name = "Courier New"
code_run.font.size = Pt(10)
inline_p.add_run(" to compute significance.")

doc.add_heading("4.2  Code Block", level=2)
code_block_lines = [
    "def fibonacci(n: int) -> int:",
    '    """Return the nth Fibonacci number."""',
    "    if n <= 1:",
    "        return n",
    "    return fibonacci(n - 1) + fibonacci(n - 2)",
    "",
    "# Example",
    "for i in range(10):",
    "    print(f'F({i}) = {fibonacci(i)}')",
]
for line in code_block_lines:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(line if line else " ")
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    # light grey background per paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)

# ══════════════════════════════════════════════
# SECTION 6 – Links & Cross-references
# ══════════════════════════════════════════════

doc.add_heading("5  Hyperlinks & Cross-references", level=1)

link_para = doc.add_paragraph("External hyperlink: ")
add_hyperlink(link_para, "https://www.w3.org/TR/WCAG21/", "W3C WCAG 2.1 Guidelines")
link_para.add_run(".")

xref_para = doc.add_paragraph("Internal cross-reference: see section ")
# Reference to the bookmark added earlier
xref_run = OxmlElement("w:r")
xref_run_rPr = OxmlElement("w:rPr")
xref_run.append(xref_run_rPr)
fld = OxmlElement("w:fldChar")
fld.set(qn("w:fldCharType"), "begin")
xref_run.append(fld)
xref_para._p.append(xref_run)
instr_run = OxmlElement("w:r")
instr_t = OxmlElement("w:instrText")
instr_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
instr_t.text = ' REF sec_typography \\h '
instr_run.append(instr_t)
xref_para._p.append(instr_run)
sep_run = OxmlElement("w:r")
sep_fc = OxmlElement("w:fldChar")
sep_fc.set(qn("w:fldCharType"), "separate")
sep_run.append(sep_fc)
xref_para._p.append(sep_run)
xref_para.add_run("1  Typography & Inline Formatting")
end_run = OxmlElement("w:r")
end_fc = OxmlElement("w:fldChar")
end_fc.set(qn("w:fldCharType"), "end")
end_run.append(end_fc)
xref_para._p.append(end_run)

# ══════════════════════════════════════════════
# SECTION 7 – Math
# ══════════════════════════════════════════════

doc.add_heading("6  Mathematical Equations", level=1)
doc.add_paragraph("The Pythagorean theorem expressed as an OMML equation:")
add_math_equation(doc, "x^2 + y^2 = z^2")

doc.add_paragraph(
    "Inline mathematical notation (plain text fallback): "
    "E = mc², where c ≈ 2.998 × 10⁸ m/s."
)

# ══════════════════════════════════════════════
# SECTION 8 – Special characters & symbols
# ══════════════════════════════════════════════

doc.add_heading("7  Special Characters & Symbols", level=1)
sym_para = doc.add_paragraph(
    "Typographic: \u2018single quotes\u2019  \u201cdouble quotes\u201d  "
    "\u2013 en-dash  \u2014 em-dash  \u2026 ellipsis  "
    "\u00a9 \u00ae \u2122  "
    "\u00bc \u00bd \u00be  "
    "\u00b0 \u00b5 \u00b6  \u2020 \u2021  "
    "\u2190 \u2192 \u2194 \u21d2  "
    "\u2200 \u2203 \u2208 \u221e \u221a \u222b"
)

doc.add_paragraph("Non-breaking space between\u00a0words.")
doc.add_paragraph("Soft hyphen: dis\u00adcrimination (hyphen should only appear at line break).")

# ══════════════════════════════════════════════
# SECTION 9 – Images / Figures
# ══════════════════════════════════════════════

doc.add_heading("8  Figures & Captions", level=1)

# We create a minimal 1×1 pixel PNG inline since no image file is guaranteed present
import io, struct, zlib

def make_minimal_png(width=200, height=120, color=(100, 149, 237)):
    """Generate a solid-colour PNG bytes object."""
    def chunk(name, data):
        c = struct.pack(">I", len(data)) + name + data
        crc = zlib.crc32(c[4:]) & 0xFFFFFFFF
        return c + struct.pack(">I", crc)

    r, g, b = color
    pixels = bytes([0] + [r, g, b] * width) * height
    compressed = zlib.compress(pixels)
    png = (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", compressed)
        + chunk(b"IEND", b"")
    )
    return io.BytesIO(png)

colours_list = [
    ((100, 149, 237), "Cornflower Blue – Figure 1"),
    ((60, 179, 113),  "Medium Sea Green – Figure 2"),
    ((255, 165, 0),   "Orange – Figure 3"),
]

for clr, label in colours_list:
    img_stream = make_minimal_png(200, 100, clr)
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(img_stream, width=Inches(3))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = caption.add_run(f"Figure: {label}")
    r.italic = True
    r.font.size = Pt(9)

# ══════════════════════════════════════════════
# SECTION 10 – Horizontal rules & page breaks
# ══════════════════════════════════════════════

doc.add_heading("9  Section Dividers", level=1)
doc.add_paragraph("Paragraph before horizontal rule.")
add_horizontal_rule(doc)
doc.add_paragraph("Paragraph after horizontal rule.")

doc.add_paragraph("A hard page break follows:").add_run()
doc.add_page_break()

# ══════════════════════════════════════════════
# SECTION 11 – Two-column section
# ══════════════════════════════════════════════

doc.add_heading("10  Multi-column Layout", level=1)

# Start a new section (continuous) with 2 columns
col_section = doc.add_section(WD_SECTION.CONTINUOUS)
cols_el = col_section._sectPr.find(qn("w:cols"))
if cols_el is None:
    cols_el = OxmlElement("w:cols")
    col_section._sectPr.append(cols_el)
cols_el.set(qn("w:num"), "2")
cols_el.set(qn("w:space"), "720")

for i in range(1, 5):
    doc.add_paragraph(
        f"Column text paragraph {i} — Lorem ipsum dolor sit amet, consectetur "
        "adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. "
        "Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris."
    )

# Revert to single column
single_section = doc.add_section(WD_SECTION.CONTINUOUS)
cols_back = single_section._sectPr.find(qn("w:cols"))
if cols_back is None:
    cols_back = OxmlElement("w:cols")
    single_section._sectPr.append(cols_back)
cols_back.set(qn("w:num"), "1")

# ══════════════════════════════════════════════
# SECTION 12 – Styles showcase
# ══════════════════════════════════════════════

doc.add_heading("11  Built-in Styles Showcase", level=1)

for style_name, sample_text in [
    ("Quote",       "This paragraph uses the 'Quote' built-in style."),
    ("Intense Quote","This paragraph uses the 'Intense Quote' built-in style."),
    ("No Spacing",   "This paragraph uses the 'No Spacing' style with no extra spacing between lines."),
]:
    try:
        p = doc.add_paragraph(sample_text, style=style_name)
    except KeyError:
        p = doc.add_paragraph(f"[style '{style_name}' not available in this template] {sample_text}")

# ══════════════════════════════════════════════
# SECTION 13 – Tracked changes placeholder
# ══════════════════════════════════════════════

doc.add_heading("12  Track Changes & Comments (markers)", level=1)
tc_para = doc.add_paragraph(
    "This paragraph contains text that would typically be subject to tracked "
    "changes review. In a full pipeline test, insertions and deletions should "
    "survive round-trips."
)
add_comment(doc, tc_para, "Review comment: consider revising this sentence.")

# ── Revision mark (insertion) ─────────────────
ins_para = doc.add_paragraph("Normal text followed by ")
ins_el = OxmlElement("w:ins")
ins_el.set(qn("w:id"), "10")
ins_el.set(qn("w:author"), "Pipeline Tester")
ins_el.set(qn("w:date"), "2026-05-20T00:00:00Z")
ins_run = OxmlElement("w:r")
ins_t = OxmlElement("w:t")
ins_t.text = "inserted text (tracked)"
ins_run.append(ins_t)
ins_el.append(ins_run)
ins_para._p.append(ins_el)
ins_para.add_run(" and more normal text.")

# ── Deletion mark ─────────────────────────────
del_para = doc.add_paragraph("Normal text followed by ")
del_el = OxmlElement("w:del")
del_el.set(qn("w:id"), "11")
del_el.set(qn("w:author"), "Pipeline Tester")
del_el.set(qn("w:date"), "2026-05-20T00:00:00Z")
del_run = OxmlElement("w:r")
del_t = OxmlElement("w:delText")
del_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
del_t.text = "deleted text (tracked)"
del_run.append(del_t)
del_el.append(del_run)
del_para._p.append(del_el)
del_para.add_run(" and continuation.")

# ══════════════════════════════════════════════
# SECTION 14 – Metadata / document properties
# ══════════════════════════════════════════════

from docx.opc.constants import RELATIONSHIP_TYPE as RT
core_props = doc.core_properties
core_props.title    = "DOCX Multi-Format Pipeline Test Document"
core_props.subject  = "Publishing pipeline validation"
core_props.author   = "Automated Generator"
core_props.keywords = "docx; pipeline; pandoc; publishing; test"
core_props.description = (
    "A specimen document covering all major DOCX features to validate "
    "multi-format conversion pipelines (PDF, HTML, EPUB, etc.)."
)
core_props.language = "en-GB"
core_props.revision = 1

# ══════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════

out_path = "c:/Wikibase/docx-pipeline-test.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
